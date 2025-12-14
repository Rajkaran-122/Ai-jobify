"""
Resume Parser using spaCy NER and BERT
Extracts structured data from resumes (96% accuracy target)
"""
import spacy
import re
from typing import Dict, List, Any, Optional
from pathlib import Path
import PyPDF2
import docx
from io import BytesIO

class ResumeParser:
    def __init__(self, model_path: Optional[str] = None):
        """
        Initialize Resume Parser
        Args:
            model_path: Path to custom trained spaCy model (if available)
        """
        # Load spaCy model (use custom trained model or default)
        try:
            if model_path and Path(model_path).exists():
                self.nlp = spacy.load(model_path)
            else:
                # For MVP, use en_core_web_sm with custom patterns
                self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Fallback if model not installed
            print("Warning: spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        # Common skills database (simplified - would be 10K+ in production)
        self.skill_keywords = {
            # Programming Languages
            "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby", "php",
            # Frameworks
            "react", "angular", "vue", "django", "flask", "fastapi", "spring", "express", "nodejs",
            # Databases
            "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "cassandra", "dynamodb",
            # Cloud & DevOps
            "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins", "gitlab ci", "github actions",
            # ML/AI
            "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "jupyter", "machine learning", "deep learning",
            # Soft Skills
            "leadership", "communication", "teamwork", "problem-solving", "agile", "scrum"
        }
        
        # Email regex pattern
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        
        # Phone regex pattern (supports US format)
        self.phone_pattern = re.compile(r'(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
        
        # Education keywords
        self.education_keywords = [
            "bachelor", "master", "phd", "doctorate", "associate", "diploma",
            "university", "college", "institute", "school"
        ]
        
        # Experience keywords
        self.experience_keywords = [
            "worked", "developed", "managed", "led", "created", "designed",
            "implemented", "built", "architected", "maintained"
        ]
    
    def extract_text_from_pdf(self, file_path_or_bytes: Any) -> str:
        """Extract text from PDF file"""
        try:
            if isinstance(file_path_or_bytes, (str, Path)):
                with open(file_path_or_bytes, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
            else:
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_path_or_bytes))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
            return text
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path_or_bytes: Any) -> str:
        """Extract text from DOCX file"""
        try:
            if isinstance(file_path_or_bytes, (str, Path)):
                doc = docx.Document(file_path_or_bytes)
            else:
                doc = docx.Document(BytesIO(file_path_or_bytes))
            
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""
    
    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract email and phone number"""
        emails = self.email_pattern.findall(text)
        phones = self.phone_pattern.findall(text)
        
        return {
            "email": emails[0] if emails else None,
            "phone": phones[0] if phones else None
        }
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from text matching skill keywords"""
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.skill_keywords:
            if skill.lower() in text_lower:
                found_skills.append(skill.title())
        
        return list(set(found_skills))  # Remove duplicates
    
    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education information"""
        if not self.nlp:
            return []
        
        doc = self.nlp(text)
        education_entries = []
        
        # Simple heuristic: look for education keywords and nearby organizations
        for sent in doc.sents:
            sent_text = sent.text.lower()
            if any(keyword in sent_text for keyword in self.education_keywords):
                # Extract organizations (universities) from this sentence
                orgs = [ent.text for ent in sent.ents if ent.label_ == "ORG"]
                if orgs:
                    education_entries.append({
                        "institution": orgs[0],
                        "description": sent.text.strip()
                    })
        
        return education_entries[:3]  # Return top 3
    
    def extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience"""
        if not self.nlp:
            return []
        
        doc = self.nlp(text)
        experience_entries = []
        
        # Extract organizations
        orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
        
        # For each organization, try to find associated job descriptions
        for org in orgs[:5]:  # Limit to top 5 organizations
            # Find sentences mentioning this org
            for sent in doc.sents:
                if org.lower() in sent.text.lower():
                    experience_entries.append({
                        "company": org,
                        "description": sent.text.strip()
                    })
                    break
        
        return experience_entries
    
    def extract_name(self, text: str) -> Optional[str]:
        """Extract person's name (usually at the top)"""
        if not self.nlp:
            # Fallback: take first line
            lines = text.strip().split('\n')
            return lines[0].strip() if lines else None
        
        doc = self.nlp(text[:500])  # Check first 500 chars
        persons = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]
        return persons[0] if persons else None
    
    def parse(self, resume_path_or_bytes: Any, file_type: str = "pdf") -> Dict[str, Any]:
        """
        Parse resume and extract structured data
        
        Args:
            resume_path_or_bytes: File path or bytes
            file_type: 'pdf' or 'docx'
        
        Returns:
            Dictionary with parsed resume data
        """
        # Extract text based on file type
        if file_type.lower() == "pdf":
            text = self.extract_text_from_pdf(resume_path_or_bytes)
        elif file_type.lower() in ["docx", "doc"]:
            text = self.extract_text_from_docx(resume_path_or_bytes)
        else:
            text = resume_path_or_bytes if isinstance(resume_path_or_bytes, str) else ""
        
        if not text:
            return {
                "personal": {},
                "education": [],
                "experience": [],
                "skills": [],
                "confidence_score": 0.0
            }
        
        # Extract all components
        contact_info = self.extract_contact_info(text)
        name = self.extract_name(text)
        skills = self.extract_skills(text)
        education = self.extract_education(text)
        experience = self.extract_experience(text)
        
        # Calculate confidence score based on extracted info
        confidence = 0.0
        if name:
            confidence += 0.2
        if contact_info.get("email"):
            confidence += 0.2
        if contact_info.get("phone"):
            confidence += 0.1
        if skills:
            confidence += 0.2
        if education:
            confidence += 0.15
        if experience:
            confidence += 0.15
        
        return {
            "personal": {
                "name": name,
                "email": contact_info.get("email"),
                "phone": contact_info.get("phone")
            },
            "education": education,
            "experience": experience,
            "skills": skills,
            "confidence_score": min(confidence, 1.0)  # Cap at 1.0
        }


# Example usage
if __name__ == "__main__":
    parser = ResumeParser()
    
    # Test with sample text
    sample_resume = """
    John Doe
    john.doe@email.com
    (555) 123-4567
    
    EDUCATION
    Bachelor of Science in Computer Science
    Stanford University, 2018-2022
    
    EXPERIENCE
    Software Engineer at Google
    Developed backend services using Python and FastAPI
    
    SKILLS
    Python, JavaScript, React, PostgreSQL, AWS, Docker
    """
    
    result = parser.parse(sample_resume, file_type="text")
    print("Parsed Resume:", result)
